import streamlit as st
import google.generativeai as genai
import json
from datetime import date
import io
import requests
import urllib.parse
import random
import re

# --- NEW LIBRARY FOR WORD DOCS ---
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="DLP Generator", layout="centered")

# --- 2. AI GENERATOR ---
def generate_lesson_content(api_key, subject, grade, quarter, content_std, perf_std, competency):
    try:
        genai.configure(api_key=api_key)
        
        # Using a standard model that is generally available
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        prompt = f"""
        You are an expert teacher. Create a JSON object for a Daily Lesson Plan (DLP).
        Subject: {subject}, Grade: {grade}, Quarter: {quarter}
        Content Standard: {content_std}
        Performance Standard: {perf_std}
        Learning Competency: {competency}

        CRITICAL INSTRUCTION: You MUST generate exactly 5 distinct assessment questions.

        Return ONLY raw JSON. No markdown formatting.
        Structure:
        {{
            "obj_1": "Cognitive objective",
            "obj_2": "Psychomotor objective",
            "obj_3": "Affective objective",
            "topic": "The main topic (include math equations like 3x^2 if needed)",
            "integration_within": "Topic within same subject",
            "integration_across": "Topic across other subject",
            "resources": {{
                "guide": "Teacher Guide reference",
                "materials": "Learner Materials reference",
                "textbook": "Textbook reference",
                "portal": "Learning Resource Portal reference",
                "other": "Other Learning Resources"
            }},
            "procedure": {{
                "review": "Review activity",
                "purpose_situation": "Real-life situation motivation description",
                "visual_prompt": "A simple 3-word visual description. Example: 'Red Apple Fruit'. NO sentences.",
                "vocabulary": "5 terms with definitions",
                "activity_main": "Main activity description",
                "explicitation": "Discussion details",
                "group_1": "Group 1 task",
                "group_2": "Group 2 task",
                "group_3": "Group 3 task",
                "generalization": "Reflection questions"
            }},
            "evaluation": {{
                "assess_q1": "Question 1 (Multiple choice or identification)",
                "assess_q2": "Question 2",
                "assess_q3": "Question 3",
                "assess_q4": "Question 4",
                "assess_q5": "Question 5",
                "assignment": "Assignment task",
                "remarks": "Remarks",
                "reflection": "Reflection"
            }}
        }}
        """
        
        response = model.generate_content(prompt)
        text = response.text
        # Clean potential markdown
        if "```json" in text:
            text = text.replace("```json", "").replace("```", "")
        return json.loads(text)
        
    except Exception as e:
        st.error(f"AI Error: {e}")
        return None

# --- 3. IMAGE FETCHER ---
def fetch_ai_image(keywords):
    if not keywords: keywords = "school_classroom"
    # Clean up the prompt
    clean_prompt = re.sub(r'[\n\r\t]', ' ', str(keywords))
    clean_prompt = re.sub(r'[^a-zA-Z0-9 ]', '', clean_prompt).strip()
    
    encoded_prompt = urllib.parse.quote(clean_prompt)
    seed = random.randint(1, 9999)
    
    # FIXED: Removed the markdown formatting from the URL string
    url = f"[https://image.pollinations.ai/prompt/](https://image.pollinations.ai/prompt/){encoded_prompt}?width=600&height=350&nologo=true&seed={seed}"
    
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=10)
        if response.status_code == 200:
            return io.BytesIO(response.content)
    except Exception:
        return None
    return None

# --- 4. DOCX HELPERS ---
def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_text(paragraph, text):
    """Parses text for ^ (superscript) and _ (subscript)."""
    if not text:
        return

    pattern = r"([^\^_]*)(([\^_])([0-9a-zA-Z\-]+))(.*)"
    current_text = str(text)
    
    if "^" not in current_text and "_" not in current_text:
        paragraph.add_run(current_text)
        return

    while True:
        match = re.match(pattern, current_text)
        if match:
            pre_text = match.group(1)
            marker = match.group(3)
            script_text = match.group(4)
            rest = match.group(5)
            
            if pre_text:
                paragraph.add_run(pre_text)
            
            run = paragraph.add_run(script_text)
            if marker == '^':
                run.font.superscript = True
            elif marker == '_':
                run.font.subscript = True
                
            current_text = rest
            if not current_text:
                break
        else:
            paragraph.add_run(current_text)
            break

def add_row(table, label, content, bold_label=True):
    """Adds a row and applies formatting to the content."""
    row_cells = table.add_row().cells
    
    # Label Column (Left)
    p_lbl = row_cells[0].paragraphs[0]
    run_lbl = p_lbl.add_run(label)
    if bold_label:
        run_lbl.bold = True
    
    # Content Column (Right)
    text_content = ""
    if isinstance(content, list):
        text_content = "\n".join([str(item) for item in content])
    else:
        text_content = str(content) if content else ""
    
    format_text(row_cells[1].paragraphs[0], text_content)

def add_section_header(table, text):
    """Adds a full-width section header with Blue background."""
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    set_cell_background(cell, "BDD7EE")

# --- 5. DOCX CREATOR ---
def create_docx(inputs, ai_data, teacher_name, principal_name, uploaded_image):
    doc = Document()
    
    # --- SETUP A4 PAGE SIZE & MARGINS ---
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # TITLE
    title = doc.add_paragraph("Daily Lesson Plan (DLP) Template")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    # --- TOP INFO TABLE ---
    table_top = doc.add_table(rows=1, cols=4)
    table_top.style = 'Table Grid'
    table_top.autofit = False
    
    table_top.columns[0].width = Inches(2.5)
    table_top.columns[1].width = Inches(1.15)
    table_top.columns[2].width = Inches(1.15)
    table_top.columns[3].width = Inches(2.5)

    def fill_cell(idx, label, value):
        cell = table_top.rows[0].cells[idx]
        p = cell.paragraphs[0]
        p.add_run(label).bold = True
        p.add_run("\n")
        format_text(p, value)

    fill_cell(0, "Subject Area:", inputs['subject'])
    fill_cell(1, "Grade Level:", inputs['grade'])
    fill_cell(2, "Quarter:", inputs['quarter'])
    fill_cell(3, "Date:", date.today().strftime('%B %d, %Y'))

    # --- MAIN CONTENT TABLE ---
    table_main = doc.add_table(rows=0, cols=2)
    table_main.style = 'Table Grid'
    table_main.autofit = False
    
    table_main.columns[0].width = Inches(2.0)
    table_main.columns[1].width = Inches(5.3)

    # Process Data
    objs = f"1. {ai_data.get('obj_1','')}\n2. {ai_data.get('obj_2','')}\n3. {ai_data.get('obj_3','')}"
    r = ai_data.get('resources', {})
    proc = ai_data.get('procedure', {})
    eval_sec = ai_data.get('evaluation', {})

    # SECTION I
    add_section_header(table_main, "I. CURRICULUM CONTENT, STANDARD AND LESSON COMPETENCIES")
    add_row(table_main, "A. Content Standard", inputs['content_std'])
    add_row(table_main, "B. Performance Standard", inputs['perf_std'])
    
    row_comp = table_main.add_row().cells
    row_comp[0].paragraphs[0].add_run("C. Learning Competencies").bold = True
    p_comp = row_comp[1].paragraphs[0]
    p_comp.add_run("Competency: ").bold = True
    format_text(p_comp, inputs['competency'])
    p_comp.add_run("\n\nObjectives:\n").bold = True
    p_comp.add_run(objs)

    add_row(table_main, "D. Content", ai_data.get('topic', ''))
    add_row(table_main, "E. Integration", f"Within: {ai_data.get('integration_within','')}\nAcross: {ai_data.get('integration_across','')}")

    # SECTION II
    add_section_header(table_main, "II. LEARNING RESOURCES")
    add_row(table_main, "Teacher Guide", r.get('guide', ''))
    add_row(table_main, "Learner’s Materials(LMs)", r.get('materials', ''))
    add_row(table_main, "Textbooks", r.get('textbook', ''))
    add_row(table_main, "Learning Resource (LR) Portal", r.get('portal', ''))
    add_row(table_main, "Other Learning Resources", r.get('other', ''))

    # SECTION III
    add_section_header(table_main, "III. TEACHING AND LEARNING PROCEDURE")
    add_row(table_main, "A. Activating Prior Knowledge", proc.get('review', ''))
    
    # --- IMAGE ROW ---
    row_img = table_main.add_row().cells
    row_img[0].paragraphs[0].add_run("B. Establishing Lesson Purpose").bold = True
    
    cell_img = row_img[1]
    format_text(cell_img.paragraphs[0], proc.get('purpose_situation', ''))
    cell_img.paragraphs[0].add_run("\n")
    
    img_data = None
    if uploaded_image:
        img_data = uploaded_image
    else:
        raw_prompt = proc.get('visual_prompt', 'school')
        img_data = fetch_ai_image(raw_prompt)
    
    if img_data:
        try:
            p_i = cell_img.add_paragraph()
            p_i.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_i = p_i.add_run()
            run_i.add_picture(img_data, width=Inches(3.5))
        except:
            cell_img.add_paragraph("[Image Error]")
    else:
        cell_img.add_paragraph("[No Image Available]")
        
    cell_img.add_paragraph(f"\nVocabulary:\n{proc.get('vocabulary','')}")

    # Rest of Section III
    add_row(table_main, "C. Developing Understanding", 
            f"Activity: {proc.get('activity_main','')}\n\nExplicitation: {proc.get('explicitation','')}\n\nGroup 1: {proc.get('group_1','')}\nGroup 2: {proc.get('group_2','')}\nGroup 3: {proc.get('group_3','')}")
    add_row(table_main, "D. Making Generalization", proc.get('generalization', ''))

    # SECTION IV
    add_section_header(table_main, "IV. EVALUATING LEARNING")
    
    # Construct list manually
    q1 = eval_sec.get('assess_q1', 'Question 1')
    q2 = eval_sec.get('assess_q2', 'Question 2')
    q3 = eval_sec.get('assess_q3', 'Question 3')
    q4 = eval_sec.get('assess_q4', 'Question 4')
    q5 = eval_sec.get('assess_q5', 'Question 5')
    
    def ensure_number(num, text):
        s_text = str(text).strip()
        if s_text.startswith(f"{num}."):
            return s_text
        return f"{num}. {s_text}"

    assessment_list = [
        ensure_number(1, q1),
        ensure_number(2, q2),
        ensure_number(3, q3),
        ensure_number(4, q4),
        ensure_number(5, q5)
    ]
    
    add_row(table_main, "A. Assessment", assessment_list)
    add_row(table_main, "B. Assignment", eval_sec.get('assignment', ''))
    add_row(table_main, "C. Remarks", eval_sec.get('remarks', ''))
    add_row(table_main, "D. Reflection", eval_sec.get('reflection', ''))

    doc.add_paragraph()

    # --- SIGNATORIES TABLE (Completed) ---
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.autofit = False
    sig_table.columns[0].width = Inches(3.65)
    sig_table.columns[1].width = Inches(3.65)
    
    # Headers
    sig_table.rows[0].cells[0].text = "Prepared by:"
    sig_table.rows[0].cells[1].text = "Noted by:"
    
    # Names
    sig_table.rows[1].cells[0].text = f"\n\n{teacher_name}\nTeacher"
    sig_table.rows[1].cells[1].text = f"\n\n{principal_name}\nPrincipal"
    
    # Bold names
    sig_table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
    sig_table.rows[1].cells[1].paragraphs[0].runs[0].bold = True

    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- 6. STREAMLIT UI (MAIN) ---
def main():
    st.title("📄 AI Lesson Plan Generator (Docx)")
    st.write("Fill in the details below to generate a DepEd-style DLP.")

    # Sidebar for API Key
    with st.sidebar:
        st.header("Settings")
        api_key = st.text_input("Google API Key", type="password")
        if not api_key:
            # Check secrets
            if "GOOGLE_API_KEY" in st.secrets:
                api_key = st.secrets["GOOGLE_API_KEY"]
            else:
                st.warning("Please enter your API Key or set it in Secrets.")

    # Form
    with st.form("dlp_form"):
        col1, col2 = st.columns(2)
        with col1:
            subject = st.text_input("Subject", "Mathematics")
            grade = st.selectbox("Grade Level", [f"Grade {i}" for i in range(1, 13)])
            quarter = st.selectbox("Quarter", ["1st Quarter", "2nd Quarter", "3rd Quarter", "4th Quarter"])
            teacher_name = st.text_input("Teacher Name", "JUAN DELA CRUZ")
        
        with col2:
            content_std = st.text_area("Content Standard", "The learner demonstrates understanding of...")
            perf_std = st.text_area("Performance Standard", "The learner is able to...")
            competency = st.text_area("Learning Competency", "Solves problems involving...")
            principal_name = st.text_input("Principal Name", "MARIA SANTOS")

        uploaded_file = st.file_uploader("Upload Image (Optional)", type=['png', 'jpg', 'jpeg'])
        submit_btn = st.form_submit_button("Generate Lesson Plan")

    # Processing
    if submit_btn and api_key:
        with st.spinner("Consulting AI... Generating Content..."):
            
            # 1. Generate Text
            ai_data = generate_lesson_content(
                api_key, subject, grade, quarter, content_std, perf_std, competency
            )

            if ai_data:
                st.success("Content Generated! Creating Word Document...")
                
                # 2. Convert Uploaded Image
                user_img = None
                if uploaded_file is not None:
                    user_img = io.BytesIO(uploaded_file.getvalue())

                # 3. Create DOCX
                docx_file = create_docx(
                    inputs={
                        "subject": subject,
                        "grade": grade,
                        "quarter": quarter,
                        "content_std": content_std,
                        "perf_std": perf_std,
                        "competency": competency
                    },
                    ai_data=ai_data,
                    teacher_name=teacher_name,
                    principal_name=principal_name,
                    uploaded_image=user_img
                )

                # 4. Download Button
                st.download_button(
                    label="📥 Download DLP (.docx)",
                    data=docx_file,
                    file_name=f"DLP_{subject}_{grade}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error("Failed to generate content. Please try again.")

if __name__ == "__main__":
    main()
